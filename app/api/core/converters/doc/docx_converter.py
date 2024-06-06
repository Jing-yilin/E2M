from api.core.converters.md_elements import (
    Header1,
    Header2,
    Header3,
    Paragraph,
    MdElement,
    MarkdownPage,
)
from api.blueprints.v1.schemas import ResponseData
from api.core.converters.base_converter import BaseConverter
from api.config import Config

from typing import List
from pathlib import Path
import os
import logging

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):

    @classmethod
    def allowed_formats(cls) -> list[str]:
        return ["doc", "docx"]

    def process_doc(self, **kwargs) -> str:
        from api.core.utils.file_utils import convert_doc_to_docx

        stem = Path(self.file_info.file_path).stem
        docx_file = os.path.join(Config.TEMP_DIR, f"{stem}.docx")
        # doc -> docx
        convert_doc_to_docx(self.file_info.file_path, docx_file)
        self.file = docx_file

        return self.process_docx(**kwargs)

    def process_docx(self, **kwargs) -> str:
        elements: List[MdElement] = []

        # method 1： default
        if Config.DOCX_CONVERTER == "default":
            logger.info(f"Converting [{self.file}] with default converter")
            from docx import Document

            doc = Document(self.file)

            """
            {
                "font size": {
                    "type": "header1",
                    "freq: 0.31,
                    "count": 100,
                    "paragraph": [1, 2, 3]
                },
                ...
            }
            """

            font_size_count = {}
            total_runs = 0

            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                for run in paragraph.runs:
                    font_size = run.font.size
                    if not font_size:
                        continue
                    total_runs += 1
                    if font_size not in font_size_count:
                        font_size_count[font_size] = {"count": 1, "paragraph": [i]}
                    else:
                        font_size_count[font_size]["count"] += 1
                        font_size_count[font_size]["paragraph"].append(i)

            if not font_size_count:
                return ResponseData(
                    status="error", error="No content found in the document"
                )

            all_font_size = set(font_size_count.keys())
            logger.debug(f"all_font_size: {all_font_size}")
            all_font_size = sorted(all_font_size, reverse=True)

            header1_font_size = 0
            header2_font_size = 0
            header3_font_size = 0

            # calculate the frequency of each font size
            for font_size in font_size_count:
                font_size_count[font_size]["freq"] = (
                    font_size_count[font_size]["count"] / total_runs
                )
                # if the frequency of a font size is greater than 0.8, it is a paragraph
                if font_size_count[font_size]["freq"] > 0.8:
                    font_size_count[font_size]["type"] = "paragraph"
                else:
                    # if the frequency of a font size is less than 0.8, it is a header
                    if len(all_font_size) >= 2 and font_size == all_font_size[0]:
                        font_size_count[font_size]["type"] = "header1"
                        header1_font_size = font_size
                    elif len(all_font_size) >= 3 and font_size == all_font_size[1]:
                        font_size_count[font_size]["type"] = "header2"
                        header2_font_size = font_size
                    elif len(all_font_size) >= 4 and font_size == all_font_size[2]:
                        font_size_count[font_size]["type"] = "header3"
                        header3_font_size = font_size

            logger.info(f"font_size_count: {font_size_count}")

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                if para.runs[0].font.size == header1_font_size:
                    elements.append(Header1(text=para.text))
                elif para.runs[0].font.size == header2_font_size:
                    elements.append(Header2(text=para.text))
                elif para.runs[0].font.size == header3_font_size:
                    elements.append(Header3(text=para.text))
                else:
                    elements.append(Paragraph(text=para.text))

        # method 2: unstructured
        elif Config.DOCX_CONVERTER == "unstructured":
            logger.info(f"Converting [{self.file}] with unstructured converter")
            from unstructured.partition.docx import partition_docx

            unstructured_elements = partition_docx(filename=self.file)
            for element in unstructured_elements:
                # todo: need a more accurate way to determine the category
                # if element.category == "Title":
                #     elements.append(Header1(element.text))
                # else:
                elements.append(Paragraph(text=element.text))
        else:
            raise ValueError(f"Invalid DOCX_CONVERTER: {Config.DOCX_CONVERTER}")

        raw = MarkdownPage.from_elements(elements).to_md()

        return raw

    def process(
        self,
        **kwargs,
    ) -> str:
        if self.file_info.file_type == "docx":
            raw = self.process_docx(**kwargs)
        elif self.file_info.file_type == "doc":
            raw = self.process_doc(**kwargs)
        else:
            raise ValueError(f"Unsupported file type: {self.file_info.file_type}")

        return raw
